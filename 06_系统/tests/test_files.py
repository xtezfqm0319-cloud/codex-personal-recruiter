from pathlib import Path

import fitz
import pytest
from docx import Document
from pypdf import PdfWriter

import recruiter.files as files
from recruiter.files import TextExtractionIncompleteError, extract_text, extract_text_with_report, sha256


def test_extract_docx_and_hash(tmp_path: Path) -> None:
    path = tmp_path / "简历.docx"
    document = Document()
    document.add_paragraph("姓名：周宁")
    document.add_paragraph("目标岗位：产品经理")
    document.save(path)
    text = extract_text(path)
    assert "周宁" in text and "产品经理" in text
    assert len(sha256(path)) == 64


def test_extract_text_pdf(tmp_path: Path) -> None:
    path = tmp_path / "文本简历.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Candidate: Alice Chen | Position: AI Product Manager | Experience: product delivery")
    document.save(path)
    document.close()
    text = extract_text(path)
    assert "Alice Chen" in text and "AI Product Manager" in text


def test_scanned_pdf_uses_local_ocr_and_reports_quality(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    path = tmp_path / "扫描简历.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with path.open("wb") as handle:
        writer.write(handle)

    monkeypatch.setattr(files, "_page_has_visual_content", lambda page: True)

    def fake_render(source: Path, pages: list[int], output: Path) -> dict[int, Path]:
        image = output / "page-0001.png"
        image.write_bytes(b"test image")
        return {1: image}

    monkeypatch.setattr(files, "_render_pdf_pages", fake_render)
    monkeypatch.setattr(
        files,
        "_ocr_images",
        lambda images: ({1: "Name: Alice Chen\nTarget Position: AI Product Manager\nExperience: enterprise AI product delivery"}, "测试OCR"),
    )

    result = extract_text_with_report(path)
    assert "Alice Chen" in result.text
    assert result.report.status == "通过"
    assert result.report.method == "本地OCR"
    assert result.report.ocr_pages == (1,)
    assert result.report.ocr_engine == "测试OCR"


def test_incomplete_ocr_pdf_requires_confirmation(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    path = tmp_path / "不完整扫描件.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with path.open("wb") as handle:
        writer.write(handle)

    monkeypatch.setattr(files, "_page_has_visual_content", lambda page: True)
    monkeypatch.setattr(files, "_render_pdf_pages", lambda source, pages, output: {1: output / "page.png"})
    monkeypatch.setattr(files, "_ocr_images", lambda images: ({1: "只有几个字"}, "测试OCR"))

    with pytest.raises(TextExtractionIncompleteError) as raised:
        extract_text_with_report(path)
    assert raised.value.report.status == "待确认"
    assert raised.value.report.unresolved_pages == (1,)


def test_image_bearing_page_compares_ocr_even_with_text_layer(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    path = tmp_path / "混合内容.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with path.open("wb") as handle:
        writer.write(handle)

    class FakePage:
        images = [object()]

        def extract_text(self) -> str:
            return "Text layer header with enough characters to pass the basic completeness threshold."

    class FakeReader:
        pages = [FakePage()]

        def __init__(self, source: str):
            pass

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)
    monkeypatch.setattr(files, "_page_has_visual_content", lambda page: True)
    monkeypatch.setattr(files, "_render_pdf_pages", lambda source, pages, output: {1: output / "page.png"})
    monkeypatch.setattr(
        files,
        "_ocr_images",
        lambda images: (
            {1: "Name: Alice Chen\nTarget Position: AI Product Manager\nExperience: complete content recovered from the embedded resume image."},
            "测试OCR",
        ),
    )

    result = extract_text_with_report(path)
    assert "complete content recovered" in result.text
    assert result.report.method == "PDF文本层 + 本地OCR"
    assert result.report.ocr_pages == (1,)
